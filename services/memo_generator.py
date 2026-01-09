# services/memo_generator.py
"""
Generador de Memorándums (RF-04, HU-03).
Genera documentos de "Solicitud de Apertura" en formato .docx.
"""

import io
import zipfile
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from datetime import datetime
from dataclasses import dataclass
from loguru import logger

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx no instalado.")

from config import get_config
from database import get_session_context, Estudiante, Proyecto, Comision


# ============================================
# CONFIGURACIÓN DE TOKENS
# ============================================

TOKENS = {
    "{{FECHA}}": "fecha_actual",
    "{{NUMERO_MEMO}}": "numero_memo",
    "{{RUN}}": "run",
    "{{NOMBRES}}": "nombres",
    "{{APELLIDOS}}": "apellidos",
    "{{NOMBRE_COMPLETO}}": "nombre_completo",
    "{{CARRERA}}": "carrera",
    "{{EMAIL}}": "email",
    "{{SEMESTRE}}": "semestre",
    "{{MODALIDAD}}": "modalidad",
    "{{TITULO_PROYECTO}}": "titulo_proyecto",
    "{{PROFESOR_GUIA}}": "profesor_guia",
    "{{CORRECTOR_1}}": "corrector_1",
    "{{CORRECTOR_2}}": "corrector_2",
}


@dataclass
class DatosMemorando:
    """Datos para generar un memorándum."""
    run: str
    nombres: str
    apellidos: str
    carrera: str
    email: Optional[str]
    semestre: Optional[str]
    modalidad: Optional[str]
    titulo_proyecto: Optional[str]
    profesor_guia: Optional[str]
    corrector_1: Optional[str]
    corrector_2: Optional[str]
    numero_memo: str = ""
    fecha: str = ""
    
    @property
    def nombre_completo(self) -> str:
        return f"{self.nombres} {self.apellidos}"
    
    def to_dict(self) -> Dict:
        return {
            "fecha_actual": self.fecha or datetime.now().strftime("%d de %B de %Y"),
            "numero_memo": self.numero_memo,
            "run": self.run,
            "nombres": self.nombres,
            "apellidos": self.apellidos,
            "nombre_completo": self.nombre_completo,
            "carrera": self.carrera,
            "email": self.email or "",
            "semestre": self.semestre or "",
            "modalidad": self.modalidad or "",
            "titulo_proyecto": self.titulo_proyecto or "",
            "profesor_guia": self.profesor_guia or "",
            "corrector_1": self.corrector_1 or "",
            "corrector_2": self.corrector_2 or "",
        }


@dataclass
class ResultadoGeneracion:
    """Resultado de la generación de memorándum."""
    run: str
    exito: bool
    ruta: Optional[str]
    nombre_archivo: str
    error: Optional[str] = None


# ============================================
# FUNCIONES DE DATOS
# ============================================

def obtener_datos_estudiante(run: str) -> Optional[DatosMemorando]:
    """Obtiene todos los datos necesarios para el memorándum."""
    try:
        with get_session_context() as session:
            estudiante = session.query(Estudiante).filter(
                Estudiante.run == run
            ).first()
            
            if not estudiante:
                logger.warning(f"Estudiante {run} no encontrado")
                return None
            
            # Buscar proyecto activo (usar estudiante_run1 que es el campo correcto)
            proyecto = session.query(Proyecto).filter(
                Proyecto.estudiante_run1 == run
            ).order_by(Proyecto.created_at.desc()).first()
            
            comision = None
            if proyecto:
                # La comisión está relacionada directamente con el proyecto
                comision = proyecto.comision
            
            return DatosMemorando(
                run=estudiante.run,
                nombres=estudiante.nombres,
                apellidos=estudiante.apellidos,
                carrera=estudiante.carrera,
                email=getattr(estudiante, 'email', None),
                semestre=proyecto.semestre if proyecto else None,
                modalidad=proyecto.modalidad_titulacion if proyecto else estudiante.modalidad,
                titulo_proyecto=proyecto.titulo if proyecto else None,
                profesor_guia=comision.profesor_guia if comision else None,
                corrector_1=comision.corrector_1 if comision else None,
                corrector_2=comision.corrector_2 if comision else None
            )
            
    except Exception as e:
        logger.error(f"Error obteniendo datos estudiante {run}: {e}", exc_info=True)
        return None


# ============================================
# GENERACIÓN DE DOCUMENTOS
# ============================================

def crear_documento_desde_plantilla(
    plantilla_path: Path,
    datos: DatosMemorando
) -> Document:
    """Crea documento reemplazando tokens en la plantilla."""
    if not HAS_DOCX:
        raise ImportError("python-docx no instalado")
    
    doc = Document(str(plantilla_path))
    datos_dict = datos.to_dict()
    
    # Reemplazar en párrafos
    for paragraph in doc.paragraphs:
        for token, campo in TOKENS.items():
            if token in paragraph.text:
                paragraph.text = paragraph.text.replace(token, datos_dict.get(campo, ""))
    
    # Reemplazar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for token, campo in TOKENS.items():
                        if token in paragraph.text:
                            paragraph.text = paragraph.text.replace(token, datos_dict.get(campo, ""))
    
    return doc


def crear_documento_por_defecto(datos: DatosMemorando) -> Document:
    """Crea un documento memorándum con formato por defecto."""
    if not HAS_DOCX:
        raise ImportError("python-docx no instalado")
    
    doc = Document()
    
    # Configurar márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)
    
    # Encabezado
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("UNIVERSIDAD DE SANTIAGO DE CHILE")
    run.bold = True
    run.font.size = Pt(14)
    
    subheader = doc.add_paragraph()
    subheader.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subheader.add_run("FACULTAD DE INGENIERÍA").bold = True
    
    doc.add_paragraph()
    
    # Título del memo
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("MEMORÁNDUM")
    run.bold = True
    run.font.size = Pt(16)
    
    doc.add_paragraph()
    
    # Datos del memo
    datos_dict = datos.to_dict()
    fecha = datetime.now().strftime("%d de %B de %Y").replace(
        "January", "Enero").replace("February", "Febrero").replace(
        "March", "Marzo").replace("April", "Abril").replace(
        "May", "Mayo").replace("June", "Junio").replace(
        "July", "Julio").replace("August", "Agosto").replace(
        "September", "Septiembre").replace("October", "Octubre").replace(
        "November", "Noviembre").replace("December", "Diciembre")
    
    doc.add_paragraph(f"FECHA: {fecha}")
    if datos.numero_memo:
        doc.add_paragraph(f"REF.: Memo N° {datos.numero_memo}")
    
    doc.add_paragraph()
    
    doc.add_paragraph("A      : REGISTRO CURRICULAR")
    doc.add_paragraph("DE    : SECRETARÍA DOCENTE - FACULTAD DE INGENIERÍA")
    doc.add_paragraph("MAT. : SOLICITUD DE APERTURA DE EXPEDIENTE DE TITULACIÓN")
    
    doc.add_paragraph()
    
    # Cuerpo
    cuerpo = doc.add_paragraph()
    cuerpo.add_run("Por medio del presente, solicito a usted la apertura del expediente de titulación para el/la estudiante:").italic = False
    
    doc.add_paragraph()
    
    # Datos del estudiante
    doc.add_paragraph(f"Nombre: {datos.nombre_completo}")
    doc.add_paragraph(f"RUN: {datos.run}")
    doc.add_paragraph(f"Carrera: {datos.carrera}")
    if datos.email:
        doc.add_paragraph(f"Email: {datos.email}")
    
    doc.add_paragraph()
    
    # Datos del proyecto
    if datos.titulo_proyecto:
        doc.add_paragraph(f"Título del Proyecto: {datos.titulo_proyecto}")
    if datos.modalidad:
        doc.add_paragraph(f"Modalidad: {datos.modalidad.replace('_', ' ').title()}")
    if datos.semestre:
        doc.add_paragraph(f"Semestre: {datos.semestre}")
    
    doc.add_paragraph()
    
    # Comisión
    if any([datos.profesor_guia, datos.corrector_1, datos.corrector_2]):
        doc.add_paragraph("Comisión Evaluadora:")
        if datos.profesor_guia:
            doc.add_paragraph(f"  - Profesor Guía: {datos.profesor_guia}")
        if datos.corrector_1:
            doc.add_paragraph(f"  - Corrector 1: {datos.corrector_1}")
        if datos.corrector_2:
            doc.add_paragraph(f"  - Corrector 2: {datos.corrector_2}")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Firma
    firma = doc.add_paragraph()
    firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    firma.add_run("_" * 40)
    
    firma2 = doc.add_paragraph()
    firma2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    firma2.add_run("Secretario/a Docente")
    
    return doc


def generar_nombre_archivo(run: str) -> str:
    """Genera nombre de archivo según estándar institucional."""
    run_limpio = run.replace(".", "").replace("-", "")
    fecha = datetime.now().strftime("%Y%m%d")
    return f"Memo_Apertura_{run_limpio}_{fecha}.docx"


# ============================================
# API PÚBLICA
# ============================================

def generar_memorandum(
    run: str,
    numero_memo: str = "",
    plantilla_path: Optional[Path] = None
) -> Tuple[bool, Optional[bytes], str]:
    """
    Genera memorándum para un estudiante.
    
    Returns:
        Tuple: (exito, bytes_documento, mensaje)
    """
    try:
        if not HAS_DOCX:
            return False, None, "python-docx no está instalado. Ejecuta: pip install python-docx"
        
        datos = obtener_datos_estudiante(run)
        if not datos:
            logger.warning(f"No se pudieron obtener datos para estudiante {run}")
            return False, None, f"Estudiante {run} no encontrado o sin datos suficientes"
        
        datos.numero_memo = numero_memo
        datos.fecha = datetime.now().strftime("%d/%m/%Y")
        
        # Buscar plantilla en múltiples ubicaciones
        config = get_config()
        plantilla = None
        
        if plantilla_path:
            # Si se especifica una ruta, usarla directamente
            plantilla = Path(plantilla_path)
        else:
            # Buscar en varias ubicaciones posibles (en orden de prioridad)
            posibles_rutas = [
                Path("C:/Users/YomiT/Downloads/Plantilla_Memo_Grado.docx"),  # Plantilla del usuario
                Path(config.paths.templates_path) / "Plantilla_Memo_Grado.docx",  # En carpeta de templates configurada
                Path(config.paths.templates_path) / "memo_template.docx",  # Nombre alternativo
                Path("./templates") / "Plantilla_Memo_Grado.docx",  # Relativo a raíz
                Path("./templates") / "memo_template.docx",  # Nombre alternativo
                Path(".") / "Plantilla_Memo_Grado.docx",  # En raíz del proyecto
                Path(".") / "memo_template.docx",  # Nombre alternativo en raíz
                Path(__file__).parent.parent / "templates" / "Plantilla_Memo_Grado.docx",  # Relativo al código
            ]
            
            for ruta in posibles_rutas:
                ruta_absoluta = ruta.resolve() if ruta.is_absolute() else Path.cwd() / ruta
                if ruta_absoluta.exists():
                    plantilla = ruta_absoluta
                    logger.info(f"✓ Plantilla encontrada en: {plantilla}")
                    break
        
        try:
            if plantilla and plantilla.exists():
                logger.info(f"Usando plantilla: {plantilla.absolute()}")
                doc = crear_documento_desde_plantilla(plantilla, datos)
            else:
                if plantilla:
                    logger.warning(f"⚠ Plantilla especificada no encontrada: {plantilla}")
                logger.warning(f"⚠ No se encontró plantilla. Buscando en: {[str(p) for p in posibles_rutas]}")
                logger.info("Usando formato por defecto (sin plantilla)")
                doc = crear_documento_por_defecto(datos)
        except Exception as e:
            logger.error(f"Error creando documento para {run}: {e}", exc_info=True)
            return False, None, f"Error creando documento: {str(e)}"
        
        # Guardar a bytes
        try:
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            bytes_doc = buffer.getvalue()
            
            if len(bytes_doc) == 0:
                return False, None, "El documento generado está vacío"
            
            nombre = generar_nombre_archivo(run)
            
            logger.info(f"Memorándum generado exitosamente para {run}: {nombre} ({len(bytes_doc)} bytes)")
            return True, bytes_doc, nombre
        except Exception as e:
            logger.error(f"Error guardando documento para {run}: {e}", exc_info=True)
            return False, None, f"Error guardando documento: {str(e)}"
        
    except Exception as e:
        logger.error(f"Error generando memorándum para {run}: {e}", exc_info=True)
        return False, None, f"Error inesperado: {str(e)}"


def generar_memorandums_masivo(
    runs: List[str],
    numero_memo_inicio: int = 1,
    callback=None
) -> Tuple[bytes, List[ResultadoGeneracion]]:
    """
    Genera memorándums para múltiples estudiantes y retorna ZIP.
    
    Args:
        runs: Lista de RUNs
        numero_memo_inicio: Número de memo inicial
        callback: Función callback(actual, total, run)
        
    Returns:
        Tuple: (bytes_zip, lista_resultados)
    """
    if not HAS_DOCX:
        logger.error("python-docx no está instalado")
        return b"", [ResultadoGeneracion(
            run=run,
            exito=False,
            ruta=None,
            nombre_archivo="",
            error="python-docx no está instalado"
        ) for run in runs]
    
    resultados = []
    documentos = []
    
    logger.info(f"Iniciando generación de {len(runs)} memorándums")
    
    for i, run in enumerate(runs):
        if callback:
            callback(i + 1, len(runs), run)
        
        numero = str(numero_memo_inicio + i).zfill(3)
        logger.debug(f"Generando memo {i+1}/{len(runs)} para {run} (número: {numero})")
        
        exito, doc_bytes, info = generar_memorandum(run, numero_memo=numero)
        
        if exito:
            documentos.append((info, doc_bytes))
            resultados.append(ResultadoGeneracion(
                run=run,
                exito=True,
                ruta=None,
                nombre_archivo=info
            ))
            logger.debug(f"✓ Memorándum generado para {run}")
        else:
            logger.warning(f"✗ Error generando memorándum para {run}: {info}")
            resultados.append(ResultadoGeneracion(
                run=run,
                exito=False,
                ruta=None,
                nombre_archivo="",
                error=info
            ))
    
    # Crear ZIP solo si hay documentos
    if len(documentos) == 0:
        logger.error("No se generó ningún memorándum. Todos fallaron.")
        return b"", resultados
    
    try:
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            for nombre, contenido in documentos:
                zf.writestr(nombre, contenido)
        
        zip_buffer.seek(0)
        zip_bytes = zip_buffer.getvalue()
        
        logger.info(f"✓ ZIP generado exitosamente: {len(documentos)}/{len(runs)} memorándums generados ({len(zip_bytes)} bytes)")
        return zip_bytes, resultados
    except Exception as e:
        logger.error(f"Error creando ZIP: {e}", exc_info=True)
        return b"", resultados
